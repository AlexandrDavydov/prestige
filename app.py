from datetime import date, datetime, timedelta
from functools import wraps
import logging
from io import BytesIO

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from flask import Flask, flash, redirect, render_template, request, send_file, session, url_for
import database as db
from database import get_all_students

app = Flask(__name__)
app.logger.setLevel(logging.INFO)
app.secret_key = "super-secret-key-l#aksjd@lakjsd"
USERNAME = "admin"
PASSWORD = "g@z"
VERSION = "1.0.0"
LESSONS_PER_PAGE = 10

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
db.init_db()



@app.route("/version", methods=["GET"])
def version():
    return "version: " + VERSION

# ===============login=====================
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        if request.form["username"] == USERNAME and request.form["password"] == PASSWORD:
            session["user"] = USERNAME
            return redirect(url_for("index"))
    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()  # очищаем сессию
    return redirect(url_for("login"))


def login_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if "user" not in session:
            return redirect(url_for("login"))
        return fn(*args, **kwargs)

    return wrapper


# ========================================
def filter_birthdays(students, target_date):
    result = []

    for student in students:
        birthday = parse_birthday(student["birthday"])

        if birthday and birthday.month == target_date.month and birthday.day == target_date.day:
            result.append(student)

    return result


def parse_birthday(value):
    if not value:
        return None

    try:
        if isinstance(value, date):
            return value  # если вдруг уже date
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return None


def get_form_value(field_name, default="", strip=True):
    value = request.form.get(field_name, default)
    return value.strip() if strip and isinstance(value, str) else value


def full_name(person):
    return f"{person['last_name']} {person['first_name']}"


def parse_student_ids(raw_student_ids):
    if not raw_student_ids:
        return []

    return [
        int(student_id)
        for student_id in raw_student_ids.split(",")
        if student_id.strip().isdigit()
    ]


def get_lesson_students(lesson):
    student_ids = set(parse_student_ids(lesson["student_ids"]))
    return [student for student in db.get_all_students() if student["id"] in student_ids]


def get_report_period():
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    if not date_from or not date_to:
        date_to = datetime.today().date()
        date_from = date_to - timedelta(days=30)
        return date_from, date_to

    return (
        datetime.strptime(date_from, "%Y-%m-%d").date(),
        datetime.strptime(date_to, "%Y-%m-%d").date(),
    )


def as_int(value, default=0):
    try:
        return int(value or default)
    except (TypeError, ValueError):
        return default


def clamp(value, minimum, maximum):
    return max(minimum, min(value, maximum))


def get_pagination_pages(current_page, total_pages):
    if total_pages <= 7:
        return list(range(1, total_pages + 1))

    visible_pages = {1, total_pages}

    for page_number in range(current_page - 1, current_page + 2):
        if 1 <= page_number <= total_pages:
            visible_pages.add(page_number)

    pages = []
    previous_page = None

    for page_number in sorted(visible_pages):
        if previous_page is not None and page_number - previous_page > 1:
            pages.append(None)
        pages.append(page_number)
        previous_page = page_number

    return pages


def paginate_items(items, page, per_page):
    total_items = len(items)
    total_pages = max(1, (total_items + per_page - 1) // per_page)
    current_page = clamp(page, 1, total_pages)
    start_index = (current_page - 1) * per_page
    end_index = start_index + per_page

    return {
        "items": items[start_index:end_index],
        "page": current_page,
        "per_page": per_page,
        "total_items": total_items,
        "total_pages": total_pages,
        "has_previous": current_page > 1,
        "has_next": current_page < total_pages,
        "pages": get_pagination_pages(current_page, total_pages),
    }


# ======= Главная =======
@app.route("/index")
@login_required
def index():
    students = get_all_students()

    today = date.today()
    yesterday = today - timedelta(days=1)
    tomorrow = today + timedelta(days=1)

    return render_template(
        "index.html",
        birthdays_yesterday=filter_birthdays(students, yesterday),
        birthdays_today=filter_birthdays(students, today),
        birthdays_tomorrow=filter_birthdays(students, tomorrow)
    )


# ======= Cards =======
@app.route("/cards")
@login_required
def cards():
    cards = db.get_all_cards()
    return render_template("cards.html", cards=cards)


@app.route("/cards/buy/<int:student_id>", methods=["GET", "POST"])
@login_required
def buy_cards(student_id):
    cards = db.get_all_cards()
    flash("Абонемент успешно добавлен!", "success")
    return render_template("buy_cards.html", cards=cards, student_id=student_id)


@app.route("/cards/purchased/", methods=["GET", "POST"])
@login_required
def purchased_cards():
    if request.method == "POST":
        db.add_card_lessons_to_student(card_id=request.form["card_id"], student_id=request.form["student_id"])
        db.store_purchased_card(card_id=request.form["card_id"], student_id=request.form["student_id"])
    return redirect(url_for("students"))


@app.route("/cards/add", methods=["GET", "POST"])
@login_required
def add_card():
    if request.method == "POST":
        db.create_card(request.form["name"], request.form["price"], request.form["lessons_count"],
                       request.form["duration"],
                       request.form["color"], "Активна"
                       )
        flash("Абонемент с названием<b> " + request.form["name"] + " </b>создан!", "success")
        return redirect(url_for("cards"))
    return render_template("add_card.html")


@app.route("/cards/edit/<int:card_id>", methods=["GET", "POST"])
@login_required
def edit_card(card_id):
    card = db.get_card_by_id(card_id)
    if request.method == "POST":
        db.update_card(card_id, request.form["name"], request.form["price"], request.form["lessons_count"],
                       request.form["duration"], request.form["color"], "Активна"
                       )
        flash("Абонемент с названием<b> " + request.form["name"] + " </b>обновлен!", "success")
        return redirect(url_for("cards"))
    return render_template("edit_card.html", card=card)


@app.route("/cards/delete/<int:card_id>")
@login_required
def delete_card(card_id):
    db.delete_card(card_id)
    flash("Абонемент удален!", "success")
    return redirect(url_for("cards"))


# ======= Students =======
@app.route("/students")
@login_required
def students():
    students = db.get_all_students()
    return render_template("students.html", students=students)


@app.route("/students/add", methods=["GET", "POST"])
@login_required
def add_student():
    if request.method == "POST":
        first_name = get_form_value("first_name")
        last_name = get_form_value("last_name")
        middle_name = get_form_value("middle_name")
        contacts = get_form_value("contacts")
        birthday = get_form_value("birthday", strip=False)
        lessons_count = as_int(get_form_value("lessons_count", 0, strip=False))
        additional_info = get_form_value("additional_info")

        # Проверяем, есть ли ученик с таким именем и фамилией
        if not db.is_student_exists(first_name, last_name):
            db.add_student(first_name, last_name, middle_name, contacts, birthday, lessons_count,
                           additional_info)
            flash(f"Ученик <b>{last_name} {first_name}</b> добавлен!", "success")
            return redirect(url_for("students"))
        else:
            flash(f"Ученик <b>{last_name} {first_name}</b> уже существует!", "error")

    # GET-запрос или если ученик уже существует
    return render_template("add_student.html")


@app.route("/students/edit/<int:student_id>", methods=["GET", "POST"])
@login_required
def edit_student(student_id):
    student = db.get_student_by_id(student_id)
    if request.method == "POST":
        db.update_student(student_id, request.form["first_name"], request.form["last_name"],
                          request.form["middle_name"],
                          request.form["contacts"], request.form["birthday"], request.form["lessons_count"],
                          request.form["additional_info"])
        flash("Данные ученика <b>" + request.form["last_name"] + " " + request.form["first_name"] + "</b> обновлены!",
              "success")
        return redirect(url_for("students"))
    return render_template("edit_student.html", student=student)


@app.route("/students/delete/<int:student_id>")
@login_required
def delete_student(student_id):
    db.delete_student(student_id)
    flash("Ученик удален!", "success")
    return redirect(url_for("students"))


# ======= Coaches =======
@app.route("/coaches")
@login_required
def coaches():
    coaches = []
    for coach_row in db.get_all_coaches():
        coach = dict(coach_row)
        coach["lessons_paid"] = as_int(coach["lessons_paid"])
        coach["student_payment"] = as_int(coach["student_payment"])
        coach["lessons_count"] = as_int(coach["lessons_count"])
        coaches.append(coach)
    return render_template("coaches.html", coaches=coaches)


@app.route("/coaches/add", methods=["GET", "POST"])
@login_required
def add_coach():
    if request.method == "POST":
        db.add_coach(request.form["first_name"], request.form["last_name"], request.form["middle_name"],
                     request.form["contacts"],
                     request.form["birthday"], request.form["lessons_count"], request.form["lessons_paid"],
                     request.form["student_payment"],
                     request.form["additional_info"])
        flash("Тренер <b>" + request.form["last_name"] + " " + request.form["first_name"] + "</b> добавлен!", "success")
        return redirect(url_for("coaches"))
    return render_template("add_coach.html")


@app.route("/coaches/edit/<int:coach_id>", methods=["GET", "POST"])
@login_required
def edit_coach(coach_id):
    coach = db.get_coach_by_id(coach_id)
    if request.method == "POST":
        db.update_coach(coach_id, request.form["first_name"], request.form["last_name"], request.form["middle_name"],
                        request.form["contacts"], request.form["birthday"], request.form["lessons_count"],
                        request.form["lessons_paid"],
                        request.form["student_payment"], request.form["additional_info"])
        flash("Данные тренера <b>" + request.form["last_name"] + " " + request.form["first_name"] + "</b> обновлены!",
              "success")
        return redirect(url_for("coaches"))
    return render_template("edit_coach.html", coach=coach)


@app.route("/coaches/delete/<int:coach_id>")
@login_required
def delete_coach(coach_id):
    db.delete_coach(coach_id)
    flash("Тренер удален!", "success")
    return redirect(url_for("coaches"))


@app.route("/coaches/money/<int:coach_id>")
@login_required
def give_money_to_coach(coach_id):
    coach = db.get_coach_by_id(coach_id)
    db.money_to_coach(coach_id)
    flash("Расчет с тренером <b>" + coach["last_name"] + " " + coach["first_name"] + "</b> произведен!", "success")
    return redirect(url_for("coaches"))


# ===== Lessons =====
@app.route("/download_docx", methods=["POST"])
def download_file_docx():
    lesson_id = request.form.get("lesson_id")
    lesson = db.get_lesson_by_id(lesson_id)
    lesson_students = get_lesson_students(lesson)
    document = Document()
    document.add_heading(lesson["lesson_name"], level=1)

    table = document.add_table(rows=len(lesson_students), cols=1)
    table.style = "Table Grid"
    table.autofit = True

    for index, student in enumerate(lesson_students):
        table.rows[index].cells[0].text = full_name(student)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/download_png", methods=["POST"])
def download_file_png():
    lesson_id = request.form.get("lesson_id")
    lesson = db.get_lesson_by_id(lesson_id)
    data = [[full_name(student)] for student in get_lesson_students(lesson)]
    cell_width = 13 * len(lesson["lesson_name"])
    cell_height = 40
    rows = len(data)
    cols = len(data[0]) if data else 1
    padding = 20
    width = cols * cell_width + padding * 2
    height = rows * cell_height + 60 + padding * 2

    img = Image.new("RGB", (width, height), color="white")
    draw = ImageDraw.Draw(img)

    # Шрифты
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 26)
        font_cell = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font_title = ImageFont.load_default()
        font_cell = ImageFont.load_default()

    # Заголовок
    title_text = lesson["lesson_name"]
    bbox = draw.textbbox((0, 0), title_text, font=font_title)
    title_width = bbox[2] - bbox[0]
    draw.text(((width - title_width) / 2, padding), title_text, fill="blue", font=font_title)

    # Таблица
    start_y = padding + 40
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            x0 = padding + j * cell_width
            y0 = start_y + i * cell_height
            x1 = x0 + cell_width
            y1 = y0 + cell_height
            draw.rectangle([x0, y0, x1, y1], outline="black")

            # Текст по центру
            bbox_cell = draw.textbbox((0, 0), cell, font=font_cell)
            cell_text_width = bbox_cell[2] - bbox_cell[0]
            cell_text_height = bbox_cell[3] - bbox_cell[1]
            draw.text(
                (
                    x0 + (cell_width - cell_text_width) / 2,
                    y0 + (cell_height - cell_text_height) / 2,
                ),
                cell,
                fill="black",
                font=font_cell,
            )

    file_stream = BytesIO()
    img.save(file_stream, format="PNG")
    file_stream.seek(0)

    return send_file(file_stream, mimetype="image/png", as_attachment=True, download_name="report.png")

@app.route("/lessons")
@login_required
def lessons():

    status_filter = request.args.get("status_filter")
    page = as_int(request.args.get("page"), default=1)

    if status_filter is None:
        status_filter = "Запланировано"

    if status_filter == "Все":
        lessons = db.get_all_lessons()
    else:
        lessons = db.get_lessons_filtered_by_status(status_filter)

    students = db.get_all_students()
    coaches = db.get_all_coaches()
    lesson_templates = db.get_all_lesson_templates()
    pagination = paginate_items(lessons, page, LESSONS_PER_PAGE)

    return render_template(
        "lessons.html",
        lessons=pagination["items"],
        students=students,
        coaches=coaches,
        lesson_templates=lesson_templates,
        status_filter=status_filter,
        pagination=pagination,
    )


@app.route("/lessons/add", methods=["GET", "POST"])
@login_required
def add_lesson():
    if request.method == "POST":
        db.add_lesson(request.form["date"], request.form["lesson_name"], request.form["coach_id"],
                      request.form.getlist("student_ids"), request.form["status"])
        flash("Занятие запланированно!", "success")
        return redirect(url_for("lessons"))
    return render_template("add_lesson.html", students=db.get_all_students(), coaches=db.get_all_coaches())


@app.route("/lessons/edit/<int:lesson_id>", methods=["GET", "POST"])
@login_required
def edit_lesson(lesson_id):
    lesson = db.get_lesson_by_id(lesson_id)
    if request.method == "POST":
        db.update_lesson(lesson_id, request.form["lesson_name"], request.form["date"], request.form["coach_id"],
                         request.form.getlist("student_ids"),
                         request.form["status"])
        flash("Занятие успешно отредактированно!", "success")
        return redirect(url_for("lessons"))

    return render_template("edit_lesson.html", lesson=lesson, students=db.get_all_students(),
                           coaches=db.get_all_coaches())


@app.route("/lessons/delete/<int:lesson_id>")
@login_required
def delete_lesson(lesson_id):
    db.delete_lesson(lesson_id)
    flash("Занятие удалено!", "success")
    return redirect(url_for("lessons"))


@app.route("/lessons/<int:lesson_id>/done", methods=["GET", "POST"])
@login_required
def lesson_done(lesson_id):
    db.decrement_lessons_from_students(lesson_id)
    db.increment_lessons_to_couch(lesson_id)
    db.close_lessons(lesson_id)
    flash("Занятие помечено как <b>состоявшееся</b>!", "success")
    return redirect(request.referrer or url_for("lessons"))


# ================================lesson template=============================
@app.route("/lesson_template/add/<int:lesson_template_id>")
@login_required
def add_lesson_from_template_route(lesson_template_id):
    db.add_lesson_from_template(lesson_template_id)
    return redirect(url_for("lessons"))


@app.route("/lesson_templates")
@login_required
def lesson_templates():
    lesson_templates = db.get_all_lesson_templates()
    students = db.get_all_students()
    coaches = db.get_all_coaches()
    return render_template("lesson_templates.html", lesson_templates=lesson_templates,
                           students=students, coaches=coaches
                           )


@app.route("/lesson_templates/add", methods=["GET", "POST"])
@login_required
def add_lesson_template():
    if request.method == "POST":
        db.add_lesson_template(request.form["template_name"], request.form["coach_id"],
                               request.form.getlist("student_ids"), )
        flash("Шаблон занятия<b> " + request.form["template_name"] + " </b>добавлен!", "success")
        return redirect(url_for("lesson_templates"))

    return render_template("add_lesson_template.html", students=db.get_all_students(),
                           coaches=db.get_all_coaches())


@app.route("/lesson_templates/edit/<int:lesson_template_id>", methods=["GET", "POST"])
@login_required
def edit_lesson_template(lesson_template_id):
    lesson_template = db.get_lesson_template_by_id(lesson_template_id)
    if request.method == "POST":
        db.update_lesson_template(lesson_template_id, request.form["template_name"], request.form["coach_id"],
                                  request.form.getlist("student_ids"),
                                  )
        flash("Шаблон занятия<b> " + request.form["template_name"] + " </b>изменен!", "success")
        return redirect(url_for("lesson_templates"))

    return render_template("edit_lesson_template.html", lesson_template=lesson_template,
                           students=db.get_all_students(), coaches=db.get_all_coaches())


@app.route("/lesson_templates/delete/<int:lesson_template_id>")
@login_required
def delete_lesson_template(lesson_template_id):
    db.delete_lesson_template(lesson_template_id)
    flash("Шаблон занятия удален!", "success")
    return redirect(url_for("lesson_templates"))


@app.template_filter("ru_date")
def ru_date(value):
    if not value:
        return ""

    for fmt in ("%Y.%m.%d", "%Y-%m-%d"):
        try:
            return datetime.strptime(value, fmt).strftime("%d.%m.%Y")
        except ValueError:
            pass

    return value  # если формат неожиданный

@app.route("/reports")
@login_required
def reports():
    return render_template("reports.html")

@app.route("/lessons_report")
@login_required
def lessons_report():
    return render_template("lessons_report.html")

def find_card_by_id(all_cards, card_id):
    for card in all_cards:
        if card["id"] == card_id:
            return card
    return None

@app.route("/cards_report")
@login_required
def cards_report():
    date_from, date_to = get_report_period()
    sold_cards_for_period = db.get_sold_cards_by_period(date_from, date_to)
    all_cards = db.get_all_cards()
    sold_cards_sum = 0

    for sold_card in sold_cards_for_period:
        card = find_card_by_id(all_cards, int(sold_card["card_id"]))
        if card:
            sold_cards_sum += int(card["price"])

    return render_template(
        "cards_report.html",
        date_from=date_from,
        date_to=date_to,
        sold_cards_for_period=sold_cards_for_period,
        sold_card_count=len(sold_cards_for_period),
        sold_cards_sum=sold_cards_sum,
    )

@app.route("/coaches_report")
@login_required
def coaches_report():
    date_from, date_to = get_report_period()
    lessons_by_period = db.get_completed_lessons_by_period(date_from, date_to)
    coach_reports = []

    for coach in db.get_all_coaches():
        coach_lessons = get_lessons_by_coach_id(lessons_by_period, coach["id"])
        coach_reports.append(
            {
                "coach_name": full_name(coach),
                "lessons_count": len(coach_lessons),
                "students_count": get_lessons_students_count(coach_lessons),
                "sum": get_lessons_sum(coach_lessons, coach["student_payment"]),
            }
        )

    totals = {
        "coach_name": "",
        "lessons_count": sum(report["lessons_count"] for report in coach_reports),
        "students_count": sum(report["students_count"] for report in coach_reports),
        "sum": sum(report["sum"] for report in coach_reports),
    }

    return render_template(
        "coaches_report.html",
        reports=[totals, coach_reports],
        date_from=date_from,
        date_to=date_to,
    )


def get_lessons_sum(lessons, coach_payment):
    total_sum = 0
    for lesson in lessons:
        if lesson["student_ids"]:
            total_sum += len(parse_student_ids(lesson["student_ids"])) * coach_payment
    return total_sum


def get_lessons_students_count(lessons):
    students_count = 0
    for lesson in lessons:
        if lesson["student_ids"]:
            students_count += len(parse_student_ids(lesson["student_ids"]))
    return students_count


def get_lessons_by_coach_id(lessons, coach_id):
    return [lesson for lesson in lessons if lesson["coach_id"] == coach_id]


@app.route("/students_report")
@login_required
def students_report():
    return render_template("students_report.html")

@app.route("/report_card_dates")
@login_required
def report_card_dates():
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    return f"<h3>Период с {date_from} по {date_to}</h3>"


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
